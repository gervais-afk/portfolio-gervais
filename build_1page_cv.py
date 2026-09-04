import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=180, bottom=140, left=320, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/><w:left w:val="none"/>'
        f'<w:bottom w:val="none"/><w:right w:val="none"/>'
        f'<w:insideH w:val="none"/><w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def set_table_zero_indent(table):
    tblPr = table._tbl.tblPr
    tblInd = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="0" w:type="dxa"/>')
    tblPr.append(tblInd)

def generate_exact_user_1page_cv():
    doc = Document()
    for s in doc.sections:
        s.page_width  = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = 0

    doc.styles['Normal'].font.name = 'Segoe UI'
    doc.styles['Normal'].font.size = Pt(8.2)

    SIDEBAR_FILL = "0F172A"
    CYAN   = RGBColor(0x38, 0xBD, 0xF8)
    ICE    = RGBColor(0xF8, 0xFA, 0xFC)
    NAVY   = RGBColor(0x0A, 0x11, 0x28)
    OCEAN  = RGBColor(0x02, 0x84, 0xC7)
    BODY   = RGBColor(0x33, 0x41, 0x55)
    MUTED  = RGBColor(0x64, 0x74, 0x8B)
    GOLD   = RGBColor(0xF5, 0xA6, 0x23)   # Accent for badge / awards
    WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_zero_indent(table)
    remove_table_borders(table)

    row = table.rows[0]
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:trHeight {nsdecls("w")} w:val="16300" w:hRule="atLeast"/>'))

    col_widths = [Inches(2.60), Inches(5.67)]

    # ══════════════ SIDEBAR (LEFT — DEEP SLATE NAVY) ══════════════
    c0 = table.cell(0, 0)
    c0.width = col_widths[0]
    set_cell_background(c0, SIDEBAR_FILL)
    set_cell_margins(c0, top=260, bottom=230, left=500, right=420)

    photo_path = r"c:\Users\HP\Desktop\portfolio-gervais\assets\images\profile_headshot_circular.jpeg"

    # ── Photo Headshot: Clean large circular photo starting at top aligned with name ──
    p_ph = c0.paragraphs[0]
    p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ph.paragraph_format.space_before = Pt(0)
    p_ph.paragraph_format.space_after  = Pt(0)
    if os.path.exists(photo_path):
        try:
            p_ph.add_run().add_picture(photo_path, width=Inches(1.60))
        except:
            r = p_ph.add_run("K G"); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = CYAN
    else:
        r = p_ph.add_run("K G"); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = CYAN

    # Badge label under photo
    p_badge = c0.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_badge.paragraph_format.space_before = Pt(3.5)
    p_badge.paragraph_format.space_after  = Pt(0)
    rb = p_badge.add_run("◈  LEAD AI ENGINEER")
    rb.font.name = 'Segoe UI'; rb.font.bold = True
    rb.font.size = Pt(7.0); rb.font.color.rgb = CYAN

    p_sub_badge = c0.add_paragraph()
    p_sub_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub_badge.paragraph_format.space_before = Pt(0)
    p_sub_badge.paragraph_format.space_after  = Pt(2)
    rs2 = p_sub_badge.add_run("Fondateur  ·  Archi Cam AI")
    rs2.font.size = Pt(6.5); rs2.font.color.rgb = MUTED

    def sb_h(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(5.5)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text.upper())
        r.font.name = 'Segoe UI'; r.font.bold = True
        r.font.size = Pt(8.4); r.font.color.rgb = CYAN
        sep = cell.add_paragraph()
        sep.paragraph_format.space_before = Pt(0.5)
        sep.paragraph_format.space_after  = Pt(2)
        rs = sep.add_run("━" * 20)
        rs.font.size = Pt(5); rs.font.color.rgb = OCEAN

    def sb_t(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1.6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(text)
        r.font.size = Pt(7.8); r.font.color.rgb = ICE

    sb_h(c0, "Contact & Profils")
    sb_t(c0, "✉  magenel85@gmail.com")
    sb_t(c0, "✆  +237 695 35 34 02")
    sb_t(c0, "⌂  Douala / Ngaoundéré, CM")
    sb_t(c0, "🌐  github.com/gervais-afk")
    sb_t(c0, "💼  linkedin.com/in/marie-gervais-koa")
    sb_t(c0, "⚡  devpost.com/magenel85")
    sb_t(c0, "🏅  Google Developer Program")

    sb_h(c0, "IA & LLM Stack")
    sb_t(c0, "Google Antigravity IDE  ■ ■ ■ ■ ■")
    sb_t(c0, "LangGraph & WikiSkills  ■ ■ ■ ■ ■")
    sb_t(c0, "Google Gemma 4 (12B)    ■ ■ ■ ■ ■")
    sb_t(c0, "Gemini 2.5 / 1.5 Pro    ■ ■ ■ ■ ■")
    sb_t(c0, "Google TabFM (Tabular)  ■ ■ ■ ■ ■")
    sb_t(c0, "Firebase Genkit         ■ ■ ■ ■ ■")
    sb_t(c0, "Neo4j GraphRAG — Agent K1")

    sb_h(c0, "Data & Graphes")
    sb_t(c0, "Neo4j / Cypher Graph    ■ ■ ■ ■ ■")
    sb_t(c0, "PostgreSQL / pgvector   ■ ■ ■ ■ ■")
    sb_t(c0, "Google BigQuery         ■ ■ ■ ■ □")
    sb_t(c0, "Supabase Realtime       ■ ■ ■ ■ □")
    sb_t(c0, "Pandas / NumPy ETL      ■ ■ ■ ■ ■")

    sb_h(c0, "Dev & MLOps")
    sb_t(c0, "Python 3.11+ / MLOps    ■ ■ ■ ■ ■")
    sb_t(c0, "FastAPI / Next.js 14    ■ ■ ■ ■ □")
    sb_t(c0, "MLflow & Data Drift     ■ ■ ■ ■ □")
    sb_t(c0, "SHAP Sentinel Audit     ■ ■ ■ ■ ■")
    sb_t(c0, "Docker & Streamlit      ■ ■ ■ ■ ■")
    sb_t(c0, "IfcOpenShell (5D BIM)   ■ ■ ■ ■ □")

    sb_h(c0, "Éthique, Sûreté & Audit")
    sb_t(c0, "◈ Éthique IA & Anti-Hallucination")
    sb_t(c0, "◈ Gouvernance Quorum 4 Yeux")
    sb_t(c0, "◈ OKF v0.2 SHA-256 No-LLM")
    sb_t(c0, "◈ Conformité EU AI Act (RSASSA)")
    sb_t(c0, "◈ Cadre AVSEC (Annexe 17 OACI)")

    sb_h(c0, "Langues")
    sb_t(c0, "Français  —  Courant / Natif")
    sb_t(c0, "Anglais   —  Bases fonctionnelles (outils & doc.")

    sb_h(c0, "Atouts Clés")
    sb_t(c0, "◈ Double compétence IA & BTP")
    sb_t(c0, "◈ Gestion de crise & Sûreté (CCAA)")
    sb_t(c0, "◈ Rigueur mathématique & Guardrails")

    # ══════════════ MAIN COLUMN (RIGHT — WHITE) ══════════════
    c1 = table.cell(0, 1)
    c1.width = col_widths[1]
    set_cell_background(c1, "FFFFFF")
    set_cell_margins(c1, top=260, bottom=230, left=280, right=300)

    # Name block — 21 pt bold
    p_nm = c1.paragraphs[0]
    p_nm.paragraph_format.space_before = Pt(0)
    p_nm.paragraph_format.space_after  = Pt(1.5)
    r = p_nm.add_run("KOA MARIE GERVAIS NELLY")
    r.font.name = 'Segoe UI'; r.font.bold = True
    r.font.size = Pt(21); r.font.color.rgb = NAVY

    p_sub = c1.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after  = Pt(1.5)
    rs = p_sub.add_run("Lead AI Engineer & Consultant IA / Data   │   Fondateur @ Archi Cam AI")
    rs.font.size = Pt(9.5); rs.font.bold = True; rs.font.color.rgb = OCEAN

    p_rule = c1.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(0)
    p_rule.paragraph_format.space_after  = Pt(3.5)
    rr = p_rule.add_run("─" * 70)
    rr.font.size = Pt(5.5); rr.font.color.rgb = OCEAN

    def mn_h(cell, title):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3.8)
        p.paragraph_format.space_after  = Pt(0.4)
        r1 = p.add_run("◈  "); r1.font.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = OCEAN
        r2 = p.add_run(title.upper())
        r2.font.name = 'Segoe UI'; r2.font.bold = True
        r2.font.size = Pt(10.2); r2.font.color.rgb = NAVY
        sep = cell.add_paragraph()
        sep.paragraph_format.space_before = Pt(0)
        sep.paragraph_format.space_after  = Pt(1.8)
        rs1 = sep.add_run("━" * 18)
        rs1.font.size = Pt(4.5); rs1.font.color.rgb = OCEAN
        rs2 = sep.add_run("─" * 44)
        rs2.font.size = Pt(4.5); rs2.font.color.rgb = CYAN

    def entry(cell, title, badge):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2.5)
        p.paragraph_format.space_after  = Pt(0.3)
        r1 = p.add_run(title); r1.font.bold = True; r1.font.size = Pt(9.3); r1.font.color.rgb = NAVY
        r2 = p.add_run(f"   —   {badge}")
        r2.font.italic = True; r2.font.size = Pt(7.9); r2.font.color.rgb = OCEAN

    def company(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0.3)
        r = p.add_run(text); r.font.size = Pt(7.9); r.font.bold = True; r.font.color.rgb = MUTED

    def bullet(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0.7)
        p.paragraph_format.line_spacing = 1.10
        p.paragraph_format.left_indent  = Inches(0.10)
        rb = p.add_run("▸  "); rb.font.bold = True; rb.font.size = Pt(7.9); rb.font.color.rgb = OCEAN
        rt = p.add_run(text); rt.font.size = Pt(7.9); rt.font.color.rgb = BODY

    def body(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1.6)
        p.paragraph_format.line_spacing = 1.13
        r = p.add_run(text); r.font.size = Pt(8.3); r.font.color.rgb = BODY

    def award(cell, title, detail):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1.0)
        p.paragraph_format.space_after  = Pt(0.6)
        p.paragraph_format.line_spacing = 1.10
        p.paragraph_format.left_indent  = Inches(0.08)
        r1 = p.add_run("◈ "); r1.font.bold = True; r1.font.size = Pt(8.2); r1.font.color.rgb = OCEAN
        r2 = p.add_run(title); r2.font.bold = True; r2.font.size = Pt(8.2); r2.font.color.rgb = NAVY
        r3 = p.add_run(f"  —  {detail}")
        r3.font.size = Pt(8.0); r3.font.color.rgb = BODY

    # ── Résumé Professionnel ──
    mn_h(c1, "Résumé Professionnel")
    body(c1, "Consultant IA & Lead AI Engineer (Google Developer Program Member), je conçois des architectures d'agents autonomes neuro-symboliques, GraphRAG (Neo4j) et pipelines MLOps souverains avec Google Antigravity. Expert en éthique de l'IA : transparence SHAP, conformité EU AI Act et modèles zéro-hallucination. Fondateur d'Archi Cam AI (candidature Google Africa Applied AI Lab 2026), j'allie rigueur mathématique, vision produit et double expertise IA / Génie Civil.")

    # ── Projets IA Majeurs ──
    mn_h(c1, "Projets IA Majeurs")

    entry(c1, "Archi Cam AI", "SaaS IA Agentique & 5D BIM")
    bullet(c1, "Candidature officielle Google Africa Applied AI Lab. Développé sous Google Antigravity (Gemma 4 12B, Gemini 2.5, BAEL 91).")
    bullet(c1, "Génération DQE Excel <45s (–99,2% temps, R²=0,9872 MLflow) et rendus 3D Imagen 3 + ControlNet.")

    entry(c1, "K1-MATHINFO (v3.0.0)", "IA Souveraine Multi-Agents, WikiSkills & Certification OKF")
    bullet(c1, "Système souverain DMI (Univ. Ngaoundéré) : 470 thèses, 19 projets M1, graphe Neo4j 1 366 nœuds.")
    bullet(c1, "6 agents LangGraph orchestrés via WikiSkills (critic anti-hallucination, Cypher optimizer), RRF k=60, SHA-256 No-LLM (100%).")

    entry(c1, "Sovereign.BI Agentic", "Business Intelligence & NL-to-SQL/Cypher")
    bullet(c1, "Moteur NL-to-SQL/Graph (PostgreSQL pgvector + Neo4j N10S, latence <5s) avec guardrails ABAC et SHAP Sentinel Audit.")

    entry(c1, "Dataset Automator & VigieSahel", "MLOps & IA Impact Agro-Climatique")
    bullet(c1, "Dataset Automator : usine MLOps Antigravity (TabFM, BigQuery DataFrames, Data Drift KS/PSI, EU AI Act).")
    bullet(c1, "VigieSahel : –35% pertes semis, anticipation épidémies +14j (XGBoost R²>94%, Supabase, MLflow).")

    # ── Parcours Professionnel ──
    mn_h(c1, "Parcours Professionnel")

    entry(c1, "Consultant IA & Data Science", "2025 – Présent")
    company(c1, "Projets Indépendants & Entreprises  │  Douala, CM")
    bullet(c1, "IA souveraines éthiques, pipelines RAG/GraphRAG Neo4j, EDA haute dimension et reporting décisionnel exécutif.")

    entry(c1, "Agent de Sûreté Aéroportuaire (AVSEC)", "2018 – Présent")
    company(c1, "CCAA — Autorité Aéronautique du Cameroun  │  Douala, CM")
    bullet(c1, "Évaluation des menaces, contrôle d'accès sécurisé, audits de conformité réglementaire ICAO Annex 17.")

    # ── Formation Académique ──
    mn_h(c1, "Formation Académique")

    entry(c1, "Master Professionnel — Intelligence Artificielle Appliquée", "2025 – 2027  [En cours]")
    company(c1, "Université de Ngaoundéré  │  Cameroun")
    bullet(c1, "ML & Stats Bayésienne, Data Engineering Neo4j, Vision & Robotique, Éthique & Cybersécurité, MLOps Souverains.")

    entry(c1, "Licence & BTS Génie Civil (Option Bâtiment)", "2015 – 2016")
    company(c1, "ISTDI / IUC Douala  │  Cameroun")
    bullet(c1, "Dimensionnement structures (BAEL 91), métrés & gestion de projets BTP — base de l'IA appliquée à l'estimation.")

    # ── Reconnaissances ──
    mn_h(c1, "Reconnaissances & Distinctions")

    p_r = c1.add_paragraph()
    p_r.paragraph_format.space_before = Pt(0)
    p_r.paragraph_format.space_after  = Pt(0)
    p_r.paragraph_format.line_spacing = 1.10
    r1 = p_r.add_run("◈ Google Africa Applied AI Lab (Accra, 2026) : ")
    r1.font.bold = True; r1.font.size = Pt(8.2); r1.font.color.rgb = NAVY
    r2 = p_r.add_run("Candidature officielle — plateforme Archi Cam AI (5D BIM + GenAI).\n")
    r2.font.size = Pt(8.0); r2.font.color.rgb = BODY
    r3 = p_r.add_run("◈ Google Cloud #AllThingsAgentic Hackathon : ")
    r3.font.bold = True; r3.font.size = Pt(8.2); r3.font.color.rgb = NAVY
    r4 = p_r.add_run("Dataset Automator v4.0 (Google Antigravity, TabFM, bigframes, WIT, MCT).")
    r4.font.size = Pt(8.0); r4.font.color.rgb = BODY

    # ── Trailing 1pt paragraph (guarantees NO extra blank page in Word) ──
    p_tail = doc.add_paragraph()
    p_tail.paragraph_format.space_before = Pt(0)
    p_tail.paragraph_format.space_after  = Pt(0)
    p_tail.paragraph_format.line_spacing = Pt(1)
    pPr = p_tail._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="20" w:lineRule="exact"/>'))
    pPr.append(parse_xml(f'<w:rPr {nsdecls("w")}><w:sz w:val="2"/><w:szCs w:val="2"/></w:rPr>'))
    r_tail = p_tail.add_run()
    r_tail.font.size = Pt(1)

    # Save DOCX
    f_docx = r"c:\Users\HP\Desktop\portfolio-gervais\KOA_MARIE_GERVAIS_NELLY_CV_FR.docx"
    try:
        doc.save(f_docx)
        print(f"Saved: {f_docx}")
    except Exception as e:
        print(f"Error saving {f_docx}: {e}")

    # PDF export and 1-page verification via Word COM
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        pdf = r"c:\Users\HP\Desktop\portfolio-gervais\KOA_MARIE_GERVAIS_NELLY_CV FR.pdf"
        if os.path.exists(f_docx):
            d = word.Documents.Open(os.path.abspath(f_docx))
            pages = d.ComputeStatistics(2)
            print(f"FR CV Page Count: {pages}")
            d.SaveAs(os.path.abspath(pdf), FileFormat=17)
            d.Close()
            print(f"Exported PDF ({pages} page): {pdf}")
        word.Quit()
        print("FR generation completed successfully!")
    except Exception as ex:
        print(f"Word COM error: {ex}")

if __name__ == "__main__":
    generate_exact_user_1page_cv()
