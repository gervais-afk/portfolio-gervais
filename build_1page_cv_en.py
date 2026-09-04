import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=180, bottom=140, left=320, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/><w:left w:val="none"/>'
        f'<w:bottom w:val="none"/><w:right w:val="none"/>'
        f'<w:insideH w:val="none"/><w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def set_table_zero_indent(table):
    tblPr = table._tbl.tblPr
    tblInd = parse_xml(f'<w:tblInd {nsdecls("w")} w:w="0" w:type="dxa"/>')
    tblPr.append(tblInd)

def generate_exact_user_1page_cv_en():
    doc = Document()
    for s in doc.sections:
        s.page_width  = Inches(8.27)
        s.page_height = Inches(11.69)
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = 0

    doc.styles['Normal'].font.name = 'Segoe UI'
    doc.styles['Normal'].font.size = Pt(8.2)

    SIDEBAR_FILL = "0F172A"
    CYAN  = RGBColor(0x38, 0xBD, 0xF8)
    ICE   = RGBColor(0xF8, 0xFA, 0xFC)
    NAVY  = RGBColor(0x0A, 0x11, 0x28)
    OCEAN = RGBColor(0x02, 0x84, 0xC7)
    BODY  = RGBColor(0x33, 0x41, 0x55)
    MUTED = RGBColor(0x64, 0x74, 0x8B)

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_zero_indent(table)
    remove_table_borders(table)

    row = table.rows[0]
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:trHeight {nsdecls("w")} w:val="16300" w:hRule="atLeast"/>'))

    col_widths = [Inches(2.60), Inches(5.67)]

    # ══════════════ SIDEBAR (LEFT — DEEP SLATE NAVY) ══════════════
    c0 = table.cell(0, 0)
    c0.width = col_widths[0]
    set_cell_background(c0, SIDEBAR_FILL)
    set_cell_margins(c0, top=260, bottom=230, left=500, right=420)

    photo_path = r"c:\Users\HP\Desktop\portfolio-gervais\assets\images\profile_headshot_circular.jpeg"

    # ── Photo Headshot: Clean large circular photo starting at top aligned with name ──
    p_ph = c0.paragraphs[0]
    p_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_ph.paragraph_format.space_before = Pt(0)
    p_ph.paragraph_format.space_after  = Pt(0)
    if os.path.exists(photo_path):
        try:
            p_ph.add_run().add_picture(photo_path, width=Inches(1.60))
        except:
            r = p_ph.add_run("K G"); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = CYAN
    else:
        r = p_ph.add_run("K G"); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = CYAN

    # Badge label under photo
    p_badge = c0.add_paragraph()
    p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_badge.paragraph_format.space_before = Pt(3.5)
    p_badge.paragraph_format.space_after  = Pt(0)
    rb = p_badge.add_run("◈  LEAD AI ENGINEER")
    rb.font.name = 'Segoe UI'; rb.font.bold = True
    rb.font.size = Pt(7.0); rb.font.color.rgb = CYAN

    p_sub_badge = c0.add_paragraph()
    p_sub_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub_badge.paragraph_format.space_before = Pt(0)
    p_sub_badge.paragraph_format.space_after  = Pt(2)
    rs2 = p_sub_badge.add_run("Founder  ·  Archi Cam AI")
    rs2.font.size = Pt(6.5); rs2.font.color.rgb = MUTED

    def sb_h(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(5.5)
        p.paragraph_format.space_after  = Pt(0)
        r = p.add_run(text.upper())
        r.font.name = 'Segoe UI'; r.font.bold = True
        r.font.size = Pt(8.4); r.font.color.rgb = CYAN
        sep = cell.add_paragraph()
        sep.paragraph_format.space_before = Pt(0.5)
        sep.paragraph_format.space_after  = Pt(2)
        rs = sep.add_run("━" * 20)
        rs.font.size = Pt(5); rs.font.color.rgb = OCEAN

    def sb_t(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1.6)
        p.paragraph_format.line_spacing = 1.08
        r = p.add_run(text)
        r.font.size = Pt(7.8); r.font.color.rgb = ICE

    sb_h(c0, "Contact & Profiles")
    sb_t(c0, "✉  magenel85@gmail.com")
    sb_t(c0, "✆  +237 695 35 34 02")
    sb_t(c0, "⌂  Douala / Ngaoundéré, CM")
    sb_t(c0, "🌐  github.com/gervais-afk")
    sb_t(c0, "💼  linkedin.com/in/marie-gervais-koa")
    sb_t(c0, "⚡  devpost.com/magenel85")
    sb_t(c0, "🏅  Google Developer Program")

    sb_h(c0, "AI & LLM Stack")
    sb_t(c0, "Google Antigravity IDE  ■ ■ ■ ■ ■")
    sb_t(c0, "LangGraph & WikiSkills  ■ ■ ■ ■ ■")
    sb_t(c0, "Google Gemma 4 (12B)    ■ ■ ■ ■ ■")
    sb_t(c0, "Gemini 2.5 / 1.5 Pro    ■ ■ ■ ■ ■")
    sb_t(c0, "Google TabFM (Tabular)  ■ ■ ■ ■ ■")
    sb_t(c0, "Firebase Genkit         ■ ■ ■ ■ ■")
    sb_t(c0, "Neo4j GraphRAG — Agent K1")

    sb_h(c0, "Data & Graphs")
    sb_t(c0, "Neo4j / Cypher Graph    ■ ■ ■ ■ ■")
    sb_t(c0, "PostgreSQL / pgvector   ■ ■ ■ ■ ■")
    sb_t(c0, "Google BigQuery         ■ ■ ■ ■ □")
    sb_t(c0, "Supabase Realtime       ■ ■ ■ ■ □")
    sb_t(c0, "Pandas / NumPy ETL      ■ ■ ■ ■ ■")

    sb_h(c0, "Dev & MLOps")
    sb_t(c0, "Python 3.11+ / MLOps    ■ ■ ■ ■ ■")
    sb_t(c0, "FastAPI / Next.js 14    ■ ■ ■ ■ □")
    sb_t(c0, "MLflow & Data Drift     ■ ■ ■ ■ □")
    sb_t(c0, "SHAP Sentinel Audit     ■ ■ ■ ■ ■")
    sb_t(c0, "Docker & Streamlit      ■ ■ ■ ■ ■")
    sb_t(c0, "IfcOpenShell (5D BIM)   ■ ■ ■ ■ □")

    sb_h(c0, "Ethics, Security & Audit")
    sb_t(c0, "◈ AI Ethics & Anti-Hallucination")
    sb_t(c0, "◈ Quorum 4-Eyes Governance")
    sb_t(c0, "◈ OKF v0.2 SHA-256 No-LLM")
    sb_t(c0, "◈ EU AI Act Compliance (RSASSA)")
    sb_t(c0, "◈ AVSEC (ICAO Annex 17)")

    sb_h(c0, "Languages")
    sb_t(c0, "French   —  Native / Fluent")
    sb_t(c0, "English  —  Functional basics (tools & doc.)")

    sb_h(c0, "Key Assets")
    sb_t(c0, "◈ Dual Competence: AI & Civil Eng.")
    sb_t(c0, "◈ Aviation Security & Crisis (CCAA)")
    sb_t(c0, "◈ Mathematical Rigor & Guardrails")

    # ══════════════ MAIN COLUMN (RIGHT — WHITE) ══════════════
    c1 = table.cell(0, 1)
    c1.width = col_widths[1]
    set_cell_background(c1, "FFFFFF")
    set_cell_margins(c1, top=240, bottom=180, left=260, right=280)

    # Name block — 21 pt bold
    p_nm = c1.paragraphs[0]
    p_nm.paragraph_format.space_before = Pt(0)
    p_nm.paragraph_format.space_after  = Pt(1.0)
    r = p_nm.add_run("KOA MARIE GERVAIS NELLY")
    r.font.name = 'Segoe UI'; r.font.bold = True
    r.font.size = Pt(21); r.font.color.rgb = NAVY

    p_sub = c1.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after  = Pt(1.0)
    rs = p_sub.add_run("Lead AI Engineer & Data Architect   │   Founder @ Archi Cam AI")
    rs.font.size = Pt(9.5); rs.font.bold = True; rs.font.color.rgb = OCEAN

    p_rule = c1.add_paragraph()
    p_rule.paragraph_format.space_before = Pt(0)
    p_rule.paragraph_format.space_after  = Pt(2.5)
    rr = p_rule.add_run("─" * 70)
    rr.font.size = Pt(5.5); rr.font.color.rgb = OCEAN

    def mn_h(cell, title):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3.2)
        p.paragraph_format.space_after  = Pt(0.2)
        r1 = p.add_run("◈  "); r1.font.bold = True; r1.font.size = Pt(8.5); r1.font.color.rgb = OCEAN
        r2 = p.add_run(title.upper())
        r2.font.name = 'Segoe UI'; r2.font.bold = True
        r2.font.size = Pt(10.0); r2.font.color.rgb = NAVY
        sep = cell.add_paragraph()
        sep.paragraph_format.space_before = Pt(0)
        sep.paragraph_format.space_after  = Pt(1.4)
        rs1 = sep.add_run("━" * 18)
        rs1.font.size = Pt(4.5); rs1.font.color.rgb = OCEAN
        rs2 = sep.add_run("─" * 44)
        rs2.font.size = Pt(4.5); rs2.font.color.rgb = CYAN

    def entry(cell, title, badge):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2.2)
        p.paragraph_format.space_after  = Pt(0.2)
        r1 = p.add_run(title); r1.font.bold = True; r1.font.size = Pt(9.3); r1.font.color.rgb = NAVY
        r2 = p.add_run(f"   —   {badge}")
        r2.font.italic = True; r2.font.size = Pt(8.0); r2.font.color.rgb = OCEAN

    def company(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0.2)
        r = p.add_run(text); r.font.size = Pt(8.0); r.font.bold = True; r.font.color.rgb = MUTED

    def bullet(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0.6)
        p.paragraph_format.line_spacing = 1.08
        p.paragraph_format.left_indent  = Inches(0.10)
        rb = p.add_run("▸  "); rb.font.bold = True; rb.font.size = Pt(8.0); rb.font.color.rgb = OCEAN
        rt = p.add_run(text); rt.font.size = Pt(8.0); rt.font.color.rgb = BODY

    def body(cell, text):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1.5)
        p.paragraph_format.line_spacing = 1.10
        r = p.add_run(text); r.font.size = Pt(8.4); r.font.color.rgb = BODY

    def award(cell, title, detail):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0.8)
        p.paragraph_format.space_after  = Pt(0.4)
        p.paragraph_format.line_spacing = 1.08
        p.paragraph_format.left_indent  = Inches(0.08)
        r1 = p.add_run("◈ "); r1.font.bold = True; r1.font.size = Pt(8.1); r1.font.color.rgb = OCEAN
        r2 = p.add_run(title); r2.font.bold = True; r2.font.size = Pt(8.1); r2.font.color.rgb = NAVY
        r3 = p.add_run(f"  —  {detail}")
        r3.font.size = Pt(8.0); r3.font.color.rgb = BODY

    # ── Executive Summary ──
    mn_h(c1, "Executive Summary")
    body(c1, "Lead AI Engineer & Data Architect (Google Developer Program Member), leveraging Google Antigravity to design neuro-symbolic multi-agent architectures, Neo4j GraphRAG, and sovereign MLOps pipelines. Expert in AI ethics: SHAP transparency, EU AI Act compliance, and deterministic zero-hallucination models. Founder & Architect of Archi Cam AI (Agentic AI SaaS & 5D BIM), bridging civil engineering methodology, mathematical rigor, and product vision.")

    # ── Flagship AI Projects ──
    mn_h(c1, "Flagship AI Projects")

    entry(c1, "Archi Cam AI", "Agentic AI SaaS & 5D BIM")
    bullet(c1, "Sovereign 5D BIM & BOQ estimation platform engineered with Google Antigravity (Gemma 4 12B, Gemini 2.5, BAEL 91).")
    bullet(c1, "Automated Excel BOQs in <45s (–99.2% time, MLflow R²=0.9872) and 3D renders via Imagen 3 + ControlNet.")
    bullet(c1, "IfcOpenShell integration for 5D BIM models: automated costing & scheduling via LangGraph agents.")

    entry(c1, "K1-MATHINFO (v3.0.0)", "Sovereign Multi-Agent AI, WikiSkills & OKF Certification")
    bullet(c1, "Sovereign DMI system (Univ. of Ngaoundéré): 470 theses, 19 M1 projects, 1,366 Neo4j nodes.")
    bullet(c1, "6 LangGraph agents orchestrated via WikiSkills (critic anti-hallucination, Cypher optimizer), RRF k=60, SHA-256 No-LLM (100%).")

    entry(c1, "Sovereign.BI Agentic", "Enterprise Security & Agentic BI")
    bullet(c1, "NL-to-SQL/Graph engine (PostgreSQL pgvector, Neo4j N10S, <5s latency) with ABAC guardrails + SHAP Sentinel.")
    bullet(c1, "Interactive executive dashboards via Streamlit & FastAPI — deployed on sovereign Docker cloud.")

    entry(c1, "Dataset Automator & VigieSahel", "MLOps Pipeline & Climate AI Impact")
    bullet(c1, "Dataset Automator: Google Antigravity (Google Cloud Hackathon) — TabFM, BigQuery DataFrames, EU AI Act.")
    bullet(c1, "VigieSahel: –35% crop loss, +14d epidemic forecast (XGBoost R²>94%, Supabase, MLflow).")

    # ── Professional Experience ──
    mn_h(c1, "Professional Experience")

    entry(c1, "AI Lead & Data Science Consultant", "March 2026 – Present")
    company(c1, "Independent Projects & Enterprises  │  Douala, CM")
    bullet(c1, "Ethical sovereign AI systems: GraphRAG Neo4j pipelines, high-dimensional EDA, multi-source RAG architectures.")
    bullet(c1, "Executive decision reporting, automated KPI dashboards, and SHAP explainability audits for African SMEs.")

    entry(c1, "Aviation Security Officer (AVSEC)", "2018 – Present")
    company(c1, "CCAA — Cameroon Civil Aviation Authority  │  Douala, CM")
    bullet(c1, "Threat assessment, secure access control, and regulatory compliance audits (ICAO Annex 17).")
    bullet(c1, "Operational crisis management: emergency team coordination, anti-intrusion protocols.")

    # ── Education ──
    mn_h(c1, "Education & Certifications")

    entry(c1, "M.Sc. in Applied Artificial Intelligence", "Dec. 2025 – 2027  [In Progress]")
    company(c1, "University of Ngaoundéré  │  Cameroon")
    bullet(c1, "ML & Bayesian Statistics, Data Engineering & Neo4j, Computer Vision & Robotics, Ethics & Cybersecurity, Production MLOps.")
    bullet(c1, "Research project: K1-MATHINFO v3 sovereign system — multi-source agent certified OKF v0.2 (SHA-256 No-LLM).")

    entry(c1, "B.Sc. in Civil Engineering (Building Option)", "2015 – 2016")
    company(c1, "ISTDI / IUC Douala  │  Cameroon")
    bullet(c1, "Structural calculations (BAEL 91), quantity surveying, construction project management — AI-applied estimation base.")

    # ── Honors ──
    mn_h(c1, "Honors & Applied AI Distinctions")

    award(c1, "CCAA Certificate of Excellence & Integrity (2023)", "Awarded by the Director General for outstanding operational performance & ethics.")
    award(c1, "Google Cloud #AllThingsAgentic Hackathon", "Dataset Automator v4.0 (Google Antigravity, TabFM, BigQuery DataFrames, WIT).")
    award(c1, "Google Developer Program · Devpost", "Active member — AI open-source contributions, hackathons & sovereign AI architect.")

    # ── Trailing 1pt paragraph (guarantees NO extra blank page in Word) ──
    p_tail = doc.add_paragraph()
    p_tail.paragraph_format.space_before = Pt(0)
    p_tail.paragraph_format.space_after  = Pt(0)
    p_tail.paragraph_format.line_spacing = Pt(1)
    pPr = p_tail._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="20" w:lineRule="exact"/>'))
    pPr.append(parse_xml(f'<w:rPr {nsdecls("w")}><w:sz w:val="2"/><w:szCs w:val="2"/></w:rPr>'))
    r_tail = p_tail.add_run()
    r_tail.font.size = Pt(1)

    # Save DOCX
    f_docx = r"c:\Users\HP\Desktop\portfolio-gervais\KOA_MARIE_GERVAIS_NELLY_CV_EN.docx"
    try:
        doc.save(f_docx)
        print(f"Saved: {f_docx}")
    except Exception as e:
        print(f"Error saving {f_docx}: {e}")

    # PDF export and 1-page verification via Word COM
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        pdf = r"c:\Users\HP\Desktop\portfolio-gervais\KOA_MARIE_GERVAIS_NELLY_CV_EN.pdf"
        if os.path.exists(f_docx):
            d = word.Documents.Open(os.path.abspath(f_docx))
            pages = d.ComputeStatistics(2)
            print(f"EN CV Page Count: {pages}")
            d.SaveAs(os.path.abspath(pdf), FileFormat=17)
            d.Close()
            print(f"Exported PDF ({pages} page): {pdf}")
        word.Quit()
        print("EN generation completed successfully!")
    except Exception as ex:
        print(f"Word COM error: {ex}")

if __name__ == "__main__":
    generate_exact_user_1page_cv_en()
