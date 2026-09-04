import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_undp_cover_letter():
    doc = docx.Document()

    # A4 Dimensions with balanced margins to fill the page
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.70)
        section.right_margin = Inches(0.70)

    NAVY  = RGBColor(0x0A, 0x11, 0x28)
    OCEAN = RGBColor(0x02, 0x84, 0xC7)
    CYAN  = RGBColor(0x02, 0x84, 0xC7)
    BODY  = RGBColor(0x1E, 0x29, 0x3B)
    MUTED = RGBColor(0x47, 0x55, 0x69)

    # ── Header / Letterhead ──
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after  = Pt(1)
    r_nm = p_name.add_run("KOA MARIE GERVAIS NELLY")
    r_nm.font.name = 'Segoe UI'
    r_nm.font.bold = True
    r_nm.font.size = Pt(19)
    r_nm.font.color.rgb = NAVY

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(3)
    r_ti = p_title.add_run("Lead AI Engineer & Data Architect   │   M.Sc. Candidate in Applied AI")
    r_ti.font.name = 'Segoe UI'
    r_ti.font.bold = True
    r_ti.font.size = Pt(10.0)
    r_ti.font.color.rgb = OCEAN

    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after  = Pt(4)
    r_co = p_contact.add_run("Douala / Ngaoundéré, Cameroon   •   magenel85@gmail.com   •   +237 695 35 34 02   •   linkedin.com/in/marie-gervais-koa")
    r_co.font.name = 'Segoe UI'
    r_co.font.size = Pt(8.8)
    r_co.font.color.rgb = MUTED

    # Dual-tone divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_before = Pt(0)
    p_div.paragraph_format.space_after  = Pt(6)
    r_d1 = p_div.add_run("━" * 25)
    r_d1.font.size = Pt(4.5); r_d1.font.color.rgb = OCEAN
    r_d2 = p_div.add_run("─" * 45)
    r_d2.font.size = Pt(4.5); r_d2.font.color.rgb = CYAN

    # Date
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after  = Pt(3.5)
    r_dt = p_date.add_run("September 4, 2026")
    r_dt.font.name = 'Segoe UI'
    r_dt.font.bold = True
    r_dt.font.size = Pt(9.2)
    r_dt.font.color.rgb = NAVY

    # Recipient
    p_rec = doc.add_paragraph()
    p_rec.paragraph_format.space_before = Pt(0)
    p_rec.paragraph_format.space_after  = Pt(4.5)
    p_rec.paragraph_format.line_spacing = 1.14
    r_r1 = p_rec.add_run("To: Selection Committee\n")
    r_r1.font.bold = True
    r_r2 = p_rec.add_run("Digital, AI and Innovation (DAI) Hub  │  Bureau for Policy and Programme Support (BPPS)\n")
    r_r3 = p_rec.add_run("United Nations Development Programme (UNDP)\n")
    r_r4 = p_rec.add_run("Job Identification: 33001 — Digital, AI and Innovation Internship: Global Call for 2026")
    r_r4.font.italic = True
    for r in [r_r1, r_r2, r_r3, r_r4]:
        r.font.name = 'Segoe UI'
        r.font.size = Pt(9.0)
        r.font.color.rgb = NAVY

    # Subject line
    p_subj = doc.add_paragraph()
    p_subj.paragraph_format.space_before = Pt(2)
    p_subj.paragraph_format.space_after  = Pt(4.5)
    r_sb = p_subj.add_run("SUBJECT: Application for Digital, AI and Innovation Internship (Home-Based) — Job ID: 33001")
    r_sb.font.name = 'Segoe UI'
    r_sb.font.bold = True
    r_sb.font.size = Pt(9.8)
    r_sb.font.color.rgb = OCEAN

    # Salutation
    p_sal = doc.add_paragraph()
    p_sal.paragraph_format.space_before = Pt(0)
    p_sal.paragraph_format.space_after  = Pt(3.5)
    r_sl = p_sal.add_run("Dear Members of the Selection Committee,")
    r_sl.font.name = 'Segoe UI'
    r_sl.font.size = Pt(9.4)
    r_sl.font.color.rgb = NAVY

    def add_p(text, before=0, after=4.5, space=1.15):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after  = Pt(after)
        p.paragraph_format.line_spacing = space
        r = p.add_run(text)
        r.font.name = 'Segoe UI'
        r.font.size = Pt(9.3)
        r.font.color.rgb = BODY
        return p

    def add_bullet(bold_head, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3.2)
        p.paragraph_format.line_spacing = 1.13
        p.paragraph_format.left_indent  = Inches(0.15)
        
        # Clean bullet
        rb = p.add_run("▪  ")
        rb.font.name = 'Segoe UI'
        rb.font.bold = True
        rb.font.size = Pt(9.3)
        rb.font.color.rgb = OCEAN

        rh = p.add_run(bold_head + ": ")
        rh.font.name = 'Segoe UI'
        rh.font.bold = True
        rh.font.size = Pt(9.3)
        rh.font.color.rgb = NAVY

        rt = p.add_run(text)
        rt.font.name = 'Segoe UI'
        rt.font.size = Pt(9.3)
        rt.font.color.rgb = BODY

    # Introduction
    add_p(
        "I am writing to submit my enthusiastic application for the Digital, AI and Innovation Internship (Job ID: 33001) "
        "within UNDP's Bureau for Policy and Programme Support (BPPS). Currently enrolled in a Master's Degree in Applied "
        "Artificial Intelligence at the University of Ngaoundéré (Dec. 2025 – 2027) with a prior Civil Engineering degree (B.Sc.), "
        "this 3-month home-based internship directly fulfills my postgraduate degree's mandatory internship curriculum requirement. "
        "It also represents a premier opportunity to place my expertise in sovereign AI systems, data engineering, and digital public goods "
        "at the service of UNDP's mission to accelerate the Sustainable Development Goals (SDGs)."
    )

    # Core Value Pillars (Bullets)
    add_bullet(
        "AI Product Build & Systems Innovation",
        "As founder of Archi Cam AI (official applicant to the Google Africa Applied AI Lab 2026), I built an agentic platform combining Gemma and Gemini LLMs with 5D BIM standards to automate construction estimation and planning in <45s (R²=0.9872). Furthermore, as an applied Master 1 research project, I engineered K1-MATHINFO v3, a sovereign multi-agent GraphRAG architecture indexing 470 academic theses with Neo4j and LangGraph, incorporating Open Knowledge Framework (OKF) zero-hallucination verification."
    )

    add_bullet(
        "Data-Driven Impact for Climate & Vulnerable Communities",
        "Dedicated to the 'Leave No One Behind' principle, I co-developed VigieSahel, an agro-climatic forecasting system combining XGBoost predictive models and Supabase to reduce crop seeding losses by 35% and anticipate epidemic risks by 14 days in the Sahel region. Additionally, through Dataset Automator (Google Cloud Hackathon), I implemented MLOps pipelines with KS/PSI data drift tracking and EU AI Act alignment frameworks."
    )

    add_bullet(
        "Institutional Rigor, Ethics & Public Service Integrity",
        "In addition to my AI capabilities, my service as an Aviation Security Officer (AVSEC) at the Cameroon Civil Aviation Authority (CCAA) has instilled uncompromising discipline and crisis management readiness. In recognition of this dedication, I was awarded the official Certificate of Excellence for Performance and Integrity (April 2023) by the Director General of CCAA. I bring this proven public-service integrity to AI governance, emphasizing explainability (SHAP Sentinel Audit), fairness, and responsible data policy."
    )

    # Conclusion & Availability
    add_p(
        "Joining the UNDP DAI Hub is an exceptional opportunity to translate cutting-edge AI research into scalable, inclusive "
        "development solutions for Country Offices. Available immediately for a 3-month remote arrangement, I am eager to contribute "
        "with technical rigor, intercultural agility, and high dedication."
    )

    add_p("Thank you very much for your consideration. I look forward to the opportunity to discuss my application with you.", after=4)

    # Sign-off
    p_sign = doc.add_paragraph()
    p_sign.paragraph_format.space_before = Pt(0)
    p_sign.paragraph_format.space_after  = Pt(1)
    r_sg = p_sign.add_run("Sincerely,")
    r_sg.font.name = 'Segoe UI'
    r_sg.font.size = Pt(9.4)
    r_sg.font.color.rgb = NAVY

    p_author = doc.add_paragraph()
    p_author.paragraph_format.space_before = Pt(0)
    p_author.paragraph_format.space_after  = Pt(0)
    r_au = p_author.add_run("KOA MARIE GERVAIS NELLY")
    r_au.font.name = 'Segoe UI'
    r_au.font.bold = True
    r_au.font.size = Pt(10.5)
    r_au.font.color.rgb = NAVY

    p_sub_au = doc.add_paragraph()
    p_sub_au.paragraph_format.space_before = Pt(0)
    p_sub_au.paragraph_format.space_after  = Pt(0)
    r_sau = p_sub_au.add_run("Lead AI Engineer & Master's Candidate in Applied AI  │  Founder @ Archi Cam AI")
    r_sau.font.name = 'Segoe UI'
    r_sau.font.size = Pt(8.8)
    r_sau.font.color.rgb = OCEAN

    # Trailing 1pt paragraph
    p_tail = doc.add_paragraph()
    p_tail.paragraph_format.space_before = Pt(0)
    p_tail.paragraph_format.space_after  = Pt(0)
    p_tail.paragraph_format.line_spacing = Pt(1)
    pPr = p_tail._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:spacing {nsdecls("w")} w:before="0" w:after="0" w:line="20" w:lineRule="exact"/>'))
    pPr.append(parse_xml(f'<w:rPr {nsdecls("w")}><w:sz w:val="2"/><w:szCs w:val="2"/></w:rPr>'))
    r_tail = p_tail.add_run()
    r_tail.font.size = Pt(1)

    # Save DOCX
    out_docx = r"c:\Users\HP\Desktop\portfolio-gervais\KOA_MARIE_GERVAIS_NELLY_COVER_LETTER_UNDP.docx"
    doc.save(out_docx)
    print(f"Saved DOCX: {out_docx}")

    # Export PDF with Word COM and check page count
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc_com = word.Documents.Open(out_docx)
        page_count = doc_com.ComputeStatistics(2)
        print(f"Cover Letter Page Count: {page_count}")
        out_pdf = r"c:\Users\HP\Desktop\portfolio-gervais\KOA_MARIE_GERVAIS_NELLY_COVER_LETTER_UNDP.pdf"
        doc_com.SaveAs(out_pdf, FileFormat=17)
        doc_com.Close()
        word.Quit()
        print(f"Exported PDF: {out_pdf}")
    except Exception as e:
        print(f"PDF export warning: {e}")

if __name__ == "__main__":
    create_undp_cover_letter()
